from flask import Flask, request, jsonify, send_file
import os
import requests
import certifi
import openai
import time
from docx import Document
from dotenv import load_dotenv
from flask_cors import CORS

# 讀取環境變數
load_dotenv()

# OpenAI API Key
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    raise ValueError("❌ 環境變數缺少 OPENAI_API_KEY")

OPENAI_API_URL = "https://api.openai.com/v1/chat/completions"

# 確保輸出目錄存在
SAVE_PATH = os.path.join(os.getcwd(), "generated_docs")
os.makedirs(SAVE_PATH, exist_ok=True)

os.environ["SSL_CERT_FILE"] = certifi.where()

app = Flask(__name__)

# ✅ 設定 CORS，允許 Netlify 前端存取
CORS(app, resources={r"/*": {"origins": "https://clever-semolina-03e5ef.netlify.app"}})

# AI 產生回答
def generate_answer(question):
    headers = {
        "Authorization": f"Bearer {OPENAI_API_KEY}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": "gpt-4",
        "messages": [
            {"role": "system", "content": (
                "你是一位專業的行銷企劃專家，擅長撰寫吸引人的產品介紹、行銷文案、品牌故事，"
                "以及讓讀者產生共鳴的創意內容。"
                "請以生動、熱情、有感染力的方式回答問題，並根據不同類型的問題調整語氣："
                "👉 產品推薦：請寫出有吸引力的產品描述，強調其獨特優勢，並附加行動呼籲（CTA）"
                "👉 廣告文案：使用強烈的標語，讓人想馬上行動"
                "👉 知識科普：請用幽默、易懂的方式解釋概念，並舉例幫助理解"
                "👉 品牌故事：營造情境，讓人對品牌產生共鳴"
                "請確保你的回答有層次感，包含：\n"
                "✅ 吸引人的標題\n"
                "✅ 清楚的內容結構（重點摘要 + 詳細解釋）\n"
                "✅ 簡單明確的行動呼籲 CTA"
            )},
            {"role": "user", "content": f"請以生動的方式回答以下問題：{question}"}
        ],
        "max_tokens": 5000,
        "temperature": 1.2,
        "top_p": 0.95,
        "frequency_penalty": 0.2,
        "presence_penalty": 0.6
    }

    try:
        response = requests.post(OPENAI_API_URL, headers=headers, json=payload, verify=certifi.where())

        print(f"🔍 OpenAI API Status Code: {response.status_code}")  
        print(f"🔍 OpenAI Response Text: {response.text}")  

        response.raise_for_status()
        data = response.json()

        answer = data.get("choices", [{}])[0].get("message", {}).get("content", "").strip()

        if not answer:
            return None  

        return answer

    except requests.exceptions.RequestException as e:
        print(f"❌ OpenAI API 錯誤: {str(e)}")
        return None

# 儲存回答為 Word 文件
def save_to_word(question, answer):
    doc = Document()
    doc.add_heading('📝 生成回答', level=1)
    
    doc.add_paragraph(f"❓ 問題：\n{question}\n")
    doc.add_paragraph(f"📢 回應：\n{answer}")

    # 設定唯一檔名
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    file_path = os.path.join(SAVE_PATH, f"簡易回答_{timestamp}.docx")

    doc.save(file_path)
    return file_path

# Flask API 路由
@app.route('/askQuestion', methods=['POST'])
def askQuestion():
    try:
        data = request.get_json()
        question = data.get('question', "").strip()
        answer = data.get('answer', "").strip()  # 接收前端傳來的 AI 回應
        download = data.get('download', False)  # 判斷是否為下載請求

        if not question:
            return jsonify({"error": "⚠️ 請輸入問題"}), 400

        if download:
            if not answer:
                return jsonify({"error": "⚠️ 無法下載，請先取得 回應"}), 400
            word_file = save_to_word(question, answer)
            return send_file(word_file, as_attachment=True, download_name=os.path.basename(word_file))

        if not answer:
            answer = generate_answer(question)
            if not answer:
                return jsonify({"error": "⚠️ 無法生成回應"}), 500

        return jsonify({"answer": answer})

    except Exception as e:
        return jsonify({"error": f"❌ 伺服器錯誤：{str(e)}"}), 500

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))  # Render 需要使用環境變數的 PORT
    app.run(host="0.0.0.0", port=port)
